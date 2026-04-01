from __future__ import annotations

import json
from datetime import datetime, UTC
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches

ROOT = Path(__file__).resolve().parents[1]
RESULTS_DIR = ROOT / "results"
REPORT_DIR = ROOT / "report"
DATA_DIR = ROOT / "data" / "ransomware_logs"
OUTPUT_PATH = REPORT_DIR / "ATMT_Wazuh_Demo_Report.docx"


def load_metrics() -> dict[str, float]:
    frame = pd.read_csv(RESULTS_DIR / "metrics_summary.csv")
    return {row["metric"]: row["value"] for _, row in frame.iterrows()}


def load_time_to_detect() -> pd.DataFrame:
    return pd.read_excel(RESULTS_DIR / "metrics_table.xlsx", sheet_name="time_to_detect")


def load_sample_sources() -> dict[str, dict[str, str]]:
    payload: dict[str, dict[str, str]] = {}
    for metadata_path in sorted(DATA_DIR.glob("*/metadata.json")):
        with metadata_path.open("r", encoding="utf-8") as handle:
            metadata = json.load(handle)
        payload[metadata_path.parent.name] = metadata
    return payload


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
    document.add_paragraph()


def format_percent(value: float) -> str:
    return f"{value:.2%}"


def format_number(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    if float(value).is_integer():
        return str(int(value))
    return f"{float(value):.2f}"


def build_document() -> Document:
    metrics = load_metrics()
    time_to_detect = load_time_to_detect()
    sample_sources = load_sample_sources()

    document = Document()
    document.core_properties.title = "ATMT Wazuh Demo Report"
    document.core_properties.subject = "Behavior-based ransomware detection demonstration"
    document.core_properties.author = "OpenAI Codex"

    document.add_heading("ATMT Wazuh Demo Report", level=0)
    document.add_paragraph("Behavior-based ransomware detection demonstration using a public Lockbit telemetry source, a Dockerized Wazuh live replay, and a safe report-generation workflow.")
    document.add_paragraph(f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}")

    document.add_heading("1. Executive Summary", level=1)
    add_paragraphs(
        document,
        [
            "This report documents a safe and repeatable ransomware-detection demo built for a professor-facing presentation.",
            "The evidence source is Splunk's public Lockbit Sysmon dataset from the attack_data repository, converted into normalized analysis rows and replay JSON artifacts.",
            "The live environment uses a Dockerized Wazuh manager that tails replayed JSON events and produces alerts in real time without executing malware.",
            f"The current sample run achieved alert-level precision of {format_percent(float(metrics.get('precision', 0.0)))} and recall of {format_percent(float(metrics.get('recall', 0.0)))}, with an average time to detect of {float(metrics.get('avg_time_to_detect_seconds', 0.0)):.2f} seconds.",
        ],
    )

    document.add_heading("2. Objectives and Scope", level=1)
    add_paragraphs(
        document,
        [
            "The objective is to demonstrate behavior-based ransomware detection that is technically credible, presentation-ready, and safe to reproduce on a host machine.",
            "The repository intentionally excludes malware acquisition or execution. The evidence set consists of a public Sysmon dataset and derived replay artifacts, while the live component is limited to replaying JSON events into a Dockerized Wazuh manager.",
        ],
    )

    document.add_heading("3. Demo Architecture", level=1)
    add_paragraphs(
        document,
        [
            "The tested classroom path uses a Dockerized Wazuh manager that tails a mounted replay file and evaluates custom JSON rules.",
            "The replay source is derived from a public Lockbit Sysmon dataset, not from live malware execution.",
            "The analysis pipeline consumes a Lockbit-derived ransomware set and a public-background benign baseline derived from the same source log, then produces report-ready charts and tables.",
        ],
    )

    document.add_heading("4. Dataset and Methodology", level=1)
    add_paragraphs(
        document,
        [
            "The source telemetry is Splunk's public Lockbit Sysmon dataset. The raw source files are stored under data/public_sources/lockbit_ransomware.",
            "Benign activity is stored under data/benign_logs and is derived from non-suspicious background rows in the public source log.",
            "Ransomware-labeled activity is stored under data/ransomware_logs/lockbit_public with alerts.csv and metadata.json derived from suspicious rows in the source dataset.",
            "A detection is counted when the Wazuh rule level is greater than or equal to 10. Time to detect is measured from the attack_start_time in each sample metadata file to the first alert at or above that threshold.",
        ],
    )

    dataset_rows = []
    for sample_name, metadata in sorted(sample_sources.items()):
        dataset_rows.append(
            [
                sample_name,
                metadata.get("family", "n/a"),
                metadata.get("source", "n/a"),
                metadata.get("attack_start_time", "n/a"),
            ]
        )
    add_table(document, ["Sample", "Family", "Source", "Attack Start"], dataset_rows)

    document.add_heading("5. Results", level=1)
    metrics_rows = [
        ["True Positives", format_number(metrics.get("true_positives"))],
        ["False Positives", format_number(metrics.get("false_positives"))],
        ["False Negatives", format_number(metrics.get("false_negatives"))],
        ["Precision", format_percent(float(metrics.get("precision", 0.0)))],
        ["Recall", format_percent(float(metrics.get("recall", 0.0)))],
        ["F1-Score", format_percent(float(metrics.get("f1_score", 0.0)))],
        ["False Positive Rate", format_percent(float(metrics.get("false_positive_rate", 0.0)))],
        ["Average Time to Detect", f"{float(metrics.get('avg_time_to_detect_seconds', 0.0)):.2f} seconds"],
    ]
    add_table(document, ["Metric", "Value"], metrics_rows)

    ttd_rows = []
    for row in time_to_detect.itertuples(index=False):
        first_detection = row.first_detection_time.isoformat() if pd.notna(row.first_detection_time) else "n/a"
        ttd_rows.append([
            str(row.sample_name),
            str(row.family),
            first_detection,
            format_number(row.time_to_detect_seconds),
        ])
    add_table(document, ["Sample", "Family", "First Detection", "TTD (s)"], ttd_rows)

    for image_name in ["detection_results.png", "top_alerts.png"]:
        image_path = RESULTS_DIR / image_name
        if image_path.exists():
            document.add_paragraph(image_name)
            document.add_picture(str(image_path), width=Inches(6.3))
            document.add_paragraph()

    document.add_heading("6. Standards Mapping", level=1)
    nist_rows = [
        ["Govern", "Defined safe demo procedure, evidence handling, and operator responsibilities."],
        ["Identify", "Tracked the public source files, derived datasets, replay artifacts, Dockerized Wazuh manager, and analysis outputs."],
        ["Protect", "Avoided live malware, treated the public dataset as read-only telemetry, and replayed JSON events instead of executing samples."],
        ["Detect", "Applied Wazuh rules, alert review, metrics computation, and time-to-detect analysis."],
        ["Respond", "Included analyst triage and explanation of triggered alerts during the demo."],
        ["Recover", "Cleared the replay file and regenerated outputs from the same known public source files."],
    ]
    add_table(document, ["NIST CSF 2.0 Function", "Demo Mapping"], nist_rows)

    iso_rows = [
        ["A.8.7", "Protection against malware", "Behavioral detection logic for ransomware-like activity."],
        ["A.8.15", "Logging", "Wazuh alert exports and replay JSON ingestion provide auditable records."],
        ["A.8.16", "Monitoring activities", "Wazuh alert review and replay monitoring are core to the demo."],
        ["A.5.24", "Incident management planning and preparation", "Repeatable workflow for evidence review and explanation."],
        ["A.8.13", "Information backup", "Generated results and sample metadata preserve reproducible evidence."],
    ]
    add_table(document, ["Control", "Name", "Demo Mapping"], iso_rows)

    document.add_heading("7. Live Demonstration Procedure", level=1)
    add_paragraphs(
        document,
        [
            "1. Start Docker Desktop and bring up the Wazuh manager from infra/docker-compose.live.yml.",
            "2. Verify that the Wazuh manager container is running and that alerts.json can be tailed.",
            "3. Present the public Lockbit source and the generated metrics from results/.",
            "4. Run simulation/replay_public_lockbit.py to append a small set of Lockbit-derived JSON events into runtime/replay/live_demo.jsonl.",
            "5. Tail /var/ossec/logs/alerts/alerts.json inside the container and point out the rule IDs and descriptions.",
        ],
    )

    document.add_heading("8. Limitations and Recommendations", level=1)
    add_paragraphs(
        document,
        [
            "The current metrics are based on a derived subset of a public Lockbit Sysmon dataset and should be presented as proof of workflow, not as a claim of production-grade coverage.",
            "The false positive rate is computed at the alert level over the public-background benign rows derived from the same source log, which is appropriate for the classroom demo but narrower than a full operational evaluation.",
            "The recommended next step is to add additional public or institution-approved telemetry sets while keeping the same safe replay model.",
        ],
    )

    document.add_heading("9. Conclusion", level=1)
    add_paragraphs(
        document,
        [
            "This implementation demonstrates a complete, safe ransomware-detection presentation workflow: public-source telemetry ingestion, metrics generation, Dockerized live replay, and standards-based reporting.",
            "It is ready for a professor demo once the Dockerized Wazuh manager is started and the replay script is run on the machine.",
        ],
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
